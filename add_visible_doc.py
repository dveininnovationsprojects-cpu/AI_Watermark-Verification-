from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
import os 
import pymupdf
from docx import Document
from docx.shared import Inches

def add_text(text,Font,opacity=100,size=30):
    if Font is None:
        Font=r'c:\WINDOWS\Fonts\BROADW.TTF'
    if opacity is None:
        opacity=30
    Font=os.path.join(r'c:\WINDOWS\Fonts',Font)
    print('opacity:',opacity)
    print('size:',size)
    font = ImageFont.truetype(Font,size)
    img = Image.new("RGBA", (400, 200), (255, 255, 255, 0))  # transparent background
    draw = ImageDraw.Draw(img)  # or default font
    draw.text((10, 80), text, fill=(255, 0, 0, opacity), font=font)

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer

def add_vis_pdf(doc,text,Font,opacity,size):
    watermark_img=add_text(text,Font,opacity*100,size)
    doc=doc.read()
    doc = pymupdf.open(stream=doc, filetype="pdf")
    for page in doc:
        rect = page.rect
        wm_rect = pymupdf.Rect(
            rect.width/4,
            rect.height/2,
            rect.width*3/4,
            rect.height/2 + 100
        )
        page.insert_image(wm_rect, stream=watermark_img)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

def add_vis_doc(doc, watermark_text,font ,opacity,size):
    doc=doc.read()
    doc = Document(BytesIO(doc))

    watermark_img = add_text(watermark_text,font,opacity*100,size)

    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(watermark_img, width=Inches(4))

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output